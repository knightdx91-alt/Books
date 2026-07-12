#!/usr/bin/env python3
"""
Build a clean, upload-ready Word (.docx) of Book One that mirrors the PDF's
structure: each front-matter piece lands on its OWN page (half-title, title,
copyright, dedication), then every chapter starts on a fresh page, then the
back matter.

Why a purpose-built docx (not "export the PDF to Word")
-------------------------------------------------------
When you paste prose into Google Docs, or let Docs convert a PDF, paragraphs
get "bunched together" because the source used soft line-breaks / spacer
paragraphs / a page geometry Docs doesn't honor. This script avoids all of
that. It writes REAL Word constructs that Google Docs, Microsoft Word, and
LibreOffice all render identically:

  * one true paragraph per paragraph (no manual line breaks inside prose),
  * hard PAGE BREAKS (WD_BREAK.PAGE) between every front-matter piece and
    before every chapter — these survive the Google Docs import intact,
  * first-line indent + justification set on the paragraph STYLE (not tabs or
    spaces), so indentation is consistent everywhere,
  * a widely-available serif (Georgia) so no font substitution reflows the
    text on upload.

Revision: reads book/genesis/saeren-chronicles/REVISION (single source of truth,
same as build_pdf.py / assemble_manuscript.py).

Output: delivery/ebook/Saeren-Chronicles-Book-One-Hazel-Academy-<REV>.docx

Run:  python3 tools/build_docx.py
(assemble the manuscript first: python3 tools/assemble_manuscript.py)
"""
import os, re, html, sys

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
try:
    REV = open(os.path.join(ROOT, "REVISION"), encoding="utf-8").read().strip()
except FileNotFoundError:
    REV = ""
_rev = f"-{REV}" if REV else ""
SRC = os.path.join(ROOT, "manuscript", f"full-manuscript{_rev}.md")

# --- mode -------------------------------------------------------------------
# Default: print / Google-Docs layout (1.5 spacing).
# --narrator: double-spaced ACX/audiobook manuscript (same content + page
#             breaks, just 2.0 line spacing and a "-narrator" filename).
NARRATOR = "--narrator" in sys.argv
_suffix = "-narrator" if NARRATOR else ""
OUT = os.path.join(ROOT, "delivery", "ebook",
                   f"Saeren-Chronicles-Book-One-Hazel-Academy{_rev}{_suffix}.docx")

BODY_FONT = "Georgia"          # safe serif present in Google Docs + Word + LibreOffice
BODY_SIZE = 12                 # pt
LINE_SPACING = 2.0 if NARRATOR else 1.5
BLACK = RGBColor(0, 0, 0)

# ----- parsing (identical contract to build_pdf.py) -------------------------

def parse(path):
    raw = open(path, encoding="utf-8").read()
    lines = raw.split("\n")
    chapters = []
    i = 0
    while i < len(lines) and not re.match(r"^CHAPTER\s+[A-Z]+\s*$", lines[i].strip()):
        i += 1
    while i < len(lines):
        m = re.match(r"^CHAPTER\s+([A-Z]+)\s*$", lines[i].strip())
        if not m:
            i += 1; continue
        num = m.group(1)
        i += 1
        while i < len(lines) and not lines[i].strip():
            i += 1
        title = lines[i].strip() if i < len(lines) else ""
        i += 1
        buf = []
        while i < len(lines) and not re.match(r"^CHAPTER\s+[A-Z]+\s*$", lines[i].strip()):
            buf.append(lines[i]); i += 1
        chapters.append((num, title, "\n".join(buf)))
    return chapters


def paras_from(text):
    blocks = re.split(r"\n\s*\n", text.strip())
    out = []
    for b in blocks:
        s = b.strip()
        if not s:
            continue
        if s in ("* * *", "***", "———", "* * * *"):
            out.append(("scene", "* * *"))
        else:
            out.append(("p", " ".join(line.strip() for line in s.split("\n"))))
    return out


# ----- inline markdown -> Word runs (**bold**, *italic*, _italic_) ----------

_TOKEN = re.compile(r"(\*\*.+?\*\*|(?<!\*)\*(?!\*).+?(?<!\*)\*(?!\*)|_.+?_)")

def add_inline(paragraph, text):
    """Append text to a paragraph, honoring **bold**/ *italic* / _italic_."""
    pos = 0
    for m in _TOKEN.finditer(text):
        if m.start() > pos:
            _run(paragraph, text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("**"):
            _run(paragraph, tok[2:-2], bold=True)
        else:
            _run(paragraph, tok[1:-1], italic=True)
        pos = m.end()
    if pos < len(text):
        _run(paragraph, text[pos:])


def _run(paragraph, txt, bold=False, italic=False):
    r = paragraph.add_run(html.unescape(txt))
    r.bold = bold
    r.italic = italic
    return r


# ----- document scaffolding -------------------------------------------------

def make_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_SIZE)
    normal.font.color.rgb = BLACK
    # ensure the East-Asian slot also points at the serif (Word quirk)
    rpr = normal.element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(a), BODY_FONT)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)


def para(doc, text="", *, align=None, italic=False, bold=False, size=None,
         indent=False, space_before=0, space_after=0, inline=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if align is not None:
        p.alignment = align
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if indent:
        pf.first_line_indent = Inches(0.3)
    if text:
        if inline:
            add_inline(p, text)
            if size or italic or bold:
                for r in p.runs:
                    if size: r.font.size = Pt(size)
                    if italic: r.italic = True
                    if bold: r.bold = True
        else:
            r = p.add_run(text)
            if size: r.font.size = Pt(size)
            r.italic = italic
            r.bold = bold
    return p


def page_break(doc):
    """A hard page break — the kind Google Docs keeps on import."""
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def blanks(doc, n):
    for _ in range(n):
        doc.add_paragraph()


# ----- build ----------------------------------------------------------------

def build():
    chapters = parse(SRC)
    doc = Document()
    make_styles(doc)

    sec = doc.sections[0]
    sec.page_height = Inches(11)
    sec.page_width = Inches(8.5)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Inches(1))

    C = WD_ALIGN_PARAGRAPH.CENTER
    J = WD_ALIGN_PARAGRAPH.JUSTIFY

    # --- half-title page ---
    blanks(doc, 9)
    para(doc, "THE SAEREN CHRONICLES", align=C, italic=True, size=16)
    page_break(doc)

    # --- title page ---
    blanks(doc, 6)
    para(doc, "THE SAEREN CHRONICLES", align=C, bold=True, size=26)
    blanks(doc, 1)
    para(doc, "Book One: Hazel Academy", align=C, italic=True, size=16)
    blanks(doc, 6)
    para(doc, "Post Peleos", align=C, size=13)
    page_break(doc)

    # --- copyright page ---
    blanks(doc, 8)
    cp = [
        "Copyright © 2026 Post Peleos",
        "",
        "All rights reserved.",
        "",
        "This is a work of fiction. Names, characters, places, and incidents "
        "are products of the author’s imagination or are used fictitiously. "
        "Any resemblance to actual persons, living or dead, events, or locales "
        "is entirely coincidental.",
        "",
        "No part of this book may be reproduced in any form or by any "
        "electronic or mechanical means, including information storage and "
        "retrieval systems, without written permission from the author, except "
        "for the use of brief quotations in a book review.",
        "",
        "ISBN 979-8-2409-9043-4",
        "",
        "First Edition",
        "",
        "Post Peleos",
    ]
    for line in cp:
        para(doc, line, align=C, size=10)
    page_break(doc)

    # --- dedication page ---
    blanks(doc, 10)
    p = doc.add_paragraph(); p.alignment = C
    r = p.add_run("For my daughter Raelynn —"); r.italic = True
    r.add_break(); r2 = p.add_run("may she have all the happiness she deserves."); r2.italic = True
    page_break(doc)

    # --- chapters ---
    for idx, (num, title, text) in enumerate(chapters):
        if idx > 0:
            page_break(doc)
        blanks(doc, 3)
        para(doc, f"CHAPTER {num}", align=C, bold=True, size=14, space_after=6)
        if title:
            para(doc, title, align=C, italic=True, size=18, inline=True, space_after=20)
        first = True
        for kind, content in paras_from(text):
            if kind == "scene":
                para(doc, "* * *", align=C, space_before=8, space_after=8)
                first = True
            else:
                para(doc, content, align=J, indent=not first, inline=True)
                first = False

    # --- back matter ---
    page_break(doc)
    para(doc, "About the Author", align=C, bold=True, size=16, space_after=18)
    para(doc, "Post Peleos writes character-driven fantasy about quiet people in "
              "loud worlds — stories that love to tease as much as they wound, and "
              "never quite stop wondering what waits out beyond the stars. *The "
              "Saeren Chronicles* is their debut series.", align=J, inline=True)

    page_break(doc)
    para(doc, "Acknowledgments", align=C, bold=True, size=16, space_after=18)
    para(doc, "To my parents, for always believing in me.", align=J)

    page_break(doc)
    para(doc, "The story continues in", align=C)
    para(doc, "THE SAEREN CHRONICLES", align=C, bold=True, size=16, space_before=6)
    para(doc, "Book Two: The Resistance", align=C, italic=True, size=13, space_after=16)
    b2 = [
        "On the fourth morning the camp put her to work, which was the kindest "
        "thing it could have done.",
        "It was Brutus who did it, the broad sunburned man from the council, who "
        "found her standing at the edge of the lanes with nothing in her hands and "
        "the look on her face of a person who has run out of things to be useful "
        "for. He did not ask her how she was, which she was grateful for, because "
        "she had heard the question a dozen times since the burning school and had "
        "not yet found a true answer that did not cost more than it was worth to "
        "give. He simply looked her up and down, the way he looked at a length of "
        "timber, and grunted, and said, “You. You’ve got two arms. Come hold this.”",
        "And running under all of it, under the washing and the games and the "
        "smell of bread, a thing she could feel the way you feel a clock in a quiet "
        "house: the count. *A week,* Jazen had said. *By this time next week I mean "
        "to be on the road to the capital.* Four days of it were already gone.",
    ]
    first = True
    for p_ in b2:
        para(doc, p_, align=J, indent=not first, inline=True)
        first = False
    para(doc, "Viridia’s war is only beginning. *The Resistance* follows her from "
              "the hidden camp to the gates of the capital — and to the choice the "
              "whole of the first book has been quietly preparing her to make.",
         align=C, italic=True, inline=True, space_before=16)

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    return OUT, len(chapters)


if __name__ == "__main__":
    if not os.path.exists(SRC):
        raise SystemExit(f"missing {SRC}\nRun: python3 tools/assemble_manuscript.py first")
    out, n = build()
    print("wrote", out)
    print(f"chapters: {n}   revision: {REV or '(none)'}")
